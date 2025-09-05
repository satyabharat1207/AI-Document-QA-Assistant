import os
import asyncio
import tempfile
import pickle
import nest_asyncio
import streamlit as st
import uuid
from io import BytesIO

from dotenv import load_dotenv
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFLoader
from langchain.schema import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# Extra
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import pandas as pd
except Exception:
    pd = None

# ---------------- Async Fix ----------------
try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())
nest_asyncio.apply()

# ---------------- Config ----------------
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

# Paths for persistent memory
CHAT_HISTORY_FILE = "chat_history.pkl"

# ---------------- Streamlit UI ----------------
st.set_page_config(page_title="AI Document Q&A", page_icon="🤖", layout="wide")

# --- CSS Customization ---
st.markdown("""
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stAppDeployButton {display: none;}
    </style>         
""",unsafe_allow_html=True)

st.markdown(
    """
    <h1 class="main-title">🤖 AI Document Q&A Assistant</h1>
    """,
    unsafe_allow_html=True
)

st.markdown(
    """
    <style>
    body {
        background: linear-gradient(270deg, #667eea, #764ba2, #ff6a00, #ee0979);
        background-size: 800% 800%;
        animation: gradientAnimation 15s ease infinite;
    }
    @keyframes gradientAnimation {
        0% {background-position:0% 50%}
        50% {background-position:100% 50%}
        100% {background-position:0% 50%}
    }

    /* Fixed, centered main title */
    .main-title {
        position: fixed;
        top: 30px;
        left: 50%;
        transform: translateX(-50%);
        z-index: 1000;
        text-align: center;
        color: white;
        font-size: 2.5rem;
        font-weight: bold;
    }
    
    /* Fixed, centered subtitle */
    .fixed-subtitle {
        position: fixed;
        top: 100px;
        left: 40%;
        transform: translateX(-50%);
        z-index: 900;
        text-align: left;
        color: white;
        font-size: 1.5rem;
        font-weight: normal;
    }
    /* Input box fixed at bottom */
    .stTextInput {
        position: fixed;
        bottom: 15px;
        left: 20%;
        right: 20%;
        z-index: 100;
    }
    .stTextInput>div>div>input {
        background-color: #ffffff;
        color: #333333;
        border: 2px solid #4a63e7;
        border-radius: 15px;
        padding: 12px 20px;
        font-size: 1rem;
        width: 100%;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.15);
        transition: all 0.3s ease;
    }
    .stTextInput>div>div>input:focus {
        outline: none;
        border-color: #764ba2;
        box-shadow: 0 0 10px rgba(118, 75, 162, 0.5);
        background-color: #f8f9fa;
    }
    .stTextInput>div>div>input:hover {
        border-color: #764ba2;
        box-shadow: 0 4px 12px rgba(118, 75, 162, 0.3);
    }
    
    /* Chat message styles using Streamlit's native components */
    .stChatMessage {
        border-radius: 12px;
        padding: 10px;
        margin: 5px 0;
        max-width: 70%;
        clear: both;
    }
    .stChatMessage.st-emotion-cache-1c7y3q8 { /* User message container */
        background-color: #4a63e7;
        color: white;
        text-align: right;
        float: right;
    }
    .stChatMessage.st-emotion-cache-1c7y3q8 > .st-emotion-cache-1c7y3q8 { /* User message content */
        color: white;
    }
    .stChatMessage.st-emotion-cache-pk30h7 { /* Bot message container */
        background-color: #f1f0f0;
        color: black;
        text-align: left;
        float: left;
    }
    .stChatMessage.st-emotion-cache-pk30h7 > .st-emotion-cache-pk30h7 { /* Bot message content */
        color: black;
    }

    </style>
    """,
    unsafe_allow_html=True
)

# ---------------- LLM & Prompts ----------------
llm = ChatGroq(groq_api_key=groq_api_key, model_name="llama-3.3-70b-versatile")

qa_prompt = ChatPromptTemplate.from_template(
    """
    Answer the question based on the provided context only.
    <context>
    {context}
    </context>
    Question: {input}
    """
)

summary_prompt = ChatPromptTemplate.from_template(
    """
    You are a document summarizer.

    First, check if the user's query is related to the given context.
    - If the query is unrelated or not found in the context, respond with:
      "❌ Your question is not provided in the documents you uploaded."
    - If the query is relevant, then summarize the context into 5 clear bullet points.

    <context>
    {context}
    </context>

    User query: {input}
    """
)

general_prompt = ChatPromptTemplate.from_template(
    """
    You are a helpful AI assistant.
    Answer the user's question clearly and concisely.

    Question: {question}
    """
)

# ---------------- Helpers: File Loading ----------------
def _save_to_temp(uploaded_file) -> str:
    suffix = os.path.splitext(uploaded_file.name)[1]
    fd, tmp_path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return tmp_path

def _load_pdf(tmp_path: str):
    return PyPDFLoader(tmp_path).load()

def _load_docx(tmp_path: str, source_name: str):
    if DocxDocument is None:
        return []
    docx = DocxDocument(tmp_path)
    text_parts = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
    return [Document(page_content="\n".join(text_parts), metadata={"source": source_name})]

def _load_csv_or_xlsx(tmp_path: str, source_name: str):
    if pd is None:
        return []
    try:
        if source_name.lower().endswith(".csv"):
            df = pd.read_csv(tmp_path, encoding="utf-8", errors="ignore")
        else:
            df = pd.read_excel(tmp_path, engine="openpyxl")
    except Exception:
        return []
    text = "\n".join([" | ".join([f"{c}: {row[c]}" for c in df.columns]) for _, row in df.iterrows()])
    return [Document(page_content=text, metadata={"source": source_name})]

def load_documents(files):
    all_docs = []
    for file in files:
        tmp_path = _save_to_temp(file)
        try:
            fname = file.name.lower()
            if fname.endswith(".pdf"):
                docs = _load_pdf(tmp_path)
            elif fname.endswith(".docx"):
                docs = _load_docx(tmp_path, file.name)
            elif fname.endswith(".csv") or fname.endswith(".xlsx"):
                docs = _load_csv_or_xlsx(tmp_path, file.name)
            else:
                docs = []
            all_docs.extend(docs)
        finally:
            os.remove(tmp_path)
    return all_docs

# ---------------- Embedding ----------------
def vector_embedding(files):
    st.session_state.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    all_docs = load_documents(files)
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(all_docs)
    st.session_state.vectors = FAISS.from_documents(chunks, st.session_state.embeddings)
    st.session_state.documents = chunks

# ---------------- Persistent Chat Memory ----------------
def save_chat_history(history):
    with open(CHAT_HISTORY_FILE, "wb") as f:
        pickle.dump(history, f)

def download_chat_history():
    history = st.session_state.chat_history
    output = "\n".join([f"{role}: {msg}" for role, msg in history])
    b = BytesIO(output.encode())
    st.download_button(label="Download Chat History", data=b, file_name="chat_history.txt", mime="text/plain")

# Initialize chat history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
    
# Initialize user input
if "user_input" not in st.session_state:
    st.session_state.user_input = ""

# ---------------- Sidebar ----------------
with st.sidebar:
    st.subheader("🧠 Chat Controls")
    if st.button("✏️ New Chat"):
        st.session_state.chat_history = []
        if os.path.exists(CHAT_HISTORY_FILE):
            os.remove(CHAT_HISTORY_FILE)
        st.toast("✅ Chat cleared, new file ready!")
    st.write("Chat history is stored in session. Download below.")
    download_chat_history()

    st.subheader("📂 Documents")
    if "uploaded_files" not in st.session_state:
        st.session_state.uploaded_files = []
    
    if "file_uploader_key" not in st.session_state:
        st.session_state.file_uploader_key = "file_uploader"
        
    uploaded_files = st.file_uploader(
        "Upload files", 
        type=["pdf", "docx", "xlsx", "csv"], 
        accept_multiple_files=True,
        key=st.session_state.file_uploader_key
    )
    
    if uploaded_files:
        st.session_state.uploaded_files = uploaded_files

    if uploaded_files and st.button("⚡ Process Documents"):
        with st.spinner("Processing..."):
            vector_embedding(st.session_state.uploaded_files)
        st.toast("✅ Documents processed successfully")
        
    
    if st.button("🗑️ Clear Documents"):
        st.session_state.pop("vectors", None)
        st.session_state.pop("documents", None)
        st.session_state.pop("embeddings", None)
        st.session_state.uploaded_files = []

        st.session_state.file_uploader_key = str(uuid.uuid4())
        st.toast("✅ Documents cleared!")
        

    st.subheader("🧩 Chat Mode")
    mode = st.radio(
        "Select a chat mode:",
        ["Q&A 📄", "Summarize 📝", "General 💬"],
        horizontal=True
    )
    st.caption("Q&A → ask from docs | Summary → quick notes | General → free chat")


# ---------------- Subtitle (fixed and centered) ----------------
st.markdown(
    """
    <h3 class="fixed-subtitle">👋 Welcome! Upload documents, choose a mode, and start chatting!</h3>
    """,
    unsafe_allow_html=True
)

# ---------------- Query Handling ----------------
def submit_query():
    if st.session_state.user_input:
        if mode == "Q&A 📄":
            if "vectors" in st.session_state and st.session_state.vectors:
                retriever = st.session_state.vectors.as_retriever()
                chain = create_stuff_documents_chain(llm, qa_prompt)
                retrieval_chain = create_retrieval_chain(retriever, chain)
                response = retrieval_chain.invoke({"input": st.session_state.user_input})
                answer = response.get("answer", "No answer.")
            else:
                answer = "⚠️ Please upload documents first for Q&A mode."

        elif mode == "Summarize 📝":
            if "vectors" in st.session_state and st.session_state.vectors:
                retriever = st.session_state.vectors.as_retriever()
                chain = create_stuff_documents_chain(llm, summary_prompt)
                retrieval_chain = create_retrieval_chain(retriever, chain)
                response = retrieval_chain.invoke({"input": st.session_state.user_input})
                answer = response.get("answer", "No summary.")
            else:
                answer = "⚠️ Please upload documents first for Summarize mode."

        elif mode == "General 💬":
            response = (general_prompt | llm).invoke({"question": st.session_state.user_input})
            answer = response.content if hasattr(response, "content") else str(response)

        # Save to history
        st.session_state.chat_history.append(("user", st.session_state.user_input))
        st.session_state.chat_history.append(("bot", answer))
        save_chat_history(st.session_state.chat_history)
        st.session_state.user_input = ""

# ---------------- Chat UI ----------------
# Display chat messages in normal order (user, then bot)
history = st.session_state.chat_history

# Group into pairs (user, bot)
pairs = []
for i in range(0, len(history), 2):
    if i + 1 < len(history):
        pairs.append((history[i], history[i + 1]))

# Show newest pair at the top
for user_msg, bot_msg in reversed(pairs):
    with st.chat_message(user_msg[0]):  # user
        st.write(user_msg[1])
    with st.chat_message(bot_msg[0]):   # bot
        st.write(bot_msg[1])


# Add JavaScript to scroll to the bottom after new messages are displayed
if st.session_state.chat_history:
    st.markdown(
        """
        <script>
        const chatContainers = window.parent.document.querySelectorAll('div[aria-label="Chat messages"]');
        if (chatContainers.length > 0) {
            const chatContainer = chatContainers[0];
            chatContainer.scrollTop = chatContainer.scrollHeight;
        }
        </script>
        """,
        unsafe_allow_html=True,
    )

# Input stays pinned
query = st.text_input("", placeholder="Ask me....", key="user_input", on_change=lambda: submit_query() if st.session_state.user_input else None)